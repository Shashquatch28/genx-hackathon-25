from dotenv import load_dotenv
from google.cloud import documentai_v1 as documentai
from google.oauth2 import service_account
from typing import Dict, Any
import tempfile
from io import BytesIO
import os
import re

load_dotenv()

# ===== define constants =====
GOOGLE_APPLICATION_CREDENTIALS = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")

PROJECT_ID = os.getenv("GCP_PROJECT_ID")
LOCATION = os.getenv("GCP_LOCATION")
PROCESSOR_ID = os.getenv("GCP_PROCESSOR_ID")

# --- Fail fast with clear messages ---
assert GOOGLE_APPLICATION_CREDENTIALS, "Missing GOOGLE_APPLICATION_CREDENTIALS in .env"
assert os.path.exists(GOOGLE_APPLICATION_CREDENTIALS), f"Service account key not found at: {GOOGLE_APPLICATION_CREDENTIALS}"
assert PROJECT_ID,  "Missing GCP_PROJECT_ID in .env"
assert LOCATION,    "Missing GCP_LOCATION in .env"
assert PROCESSOR_ID,"Missing GCP_PROCESSOR_ID in .env"

PDF_MIME = "application/pdf"
TXT_MIME = "text/plain"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
SUPPORTED_IMAGE_MIMES = {"image/jpeg", "image/png", "image/tiff", "image/gif"}

# ===== GCP Client Setup =====
def _processor_name() -> str:
    return f"projects/{PROJECT_ID}/locations/{LOCATION}/processors/{PROCESSOR_ID}"

def _client() -> documentai.DocumentProcessorServiceClient:
    """
    Build a client that uses either:
      - LOCATION as a short region code (e.g., 'eu' -> 'eu-documentai.googleapis.com')
      - or LOCATION already set to a full endpoint 'eu-documentai.googleapis.com'
    """
    credentials = service_account.Credentials.from_service_account_file(GOOGLE_APPLICATION_CREDENTIALS)

    # If LOCATION already looks like an endpoint, use it directly
    if "documentai.googleapis.com" in LOCATION:
        api_endpoint = LOCATION
    else:
        api_endpoint = f"{LOCATION}-documentai.googleapis.com"

    client_options = {"api_endpoint": api_endpoint}
    return documentai.DocumentProcessorServiceClient(
        credentials=credentials,
        client_options=client_options
    )

# ===== Process and parse input =====
def _process_with_layout(file_bytes: bytes, mime_type: str) -> documentai.Document:
    client = _client()
    name = _processor_name()

    process_options = documentai.ProcessOptions(
        layout_config=documentai.ProcessOptions.LayoutConfig(
            chunking_config=documentai.ProcessOptions.LayoutConfig.ChunkingConfig(
                chunk_size=1000,
                include_ancestor_headings=True
            )
        )
    )

    request = documentai.ProcessRequest(
        name=name,
        raw_document=documentai.RawDocument(content=file_bytes, mime_type=mime_type),
        process_options=process_options,
    )

    try:
        result = client.process_document(request=request)
    except Exception as e:
        # Surface a clearer error for debugging
        raise RuntimeError(f"Document AI processing failed: {e}") from e

    return result.document

# ===== Convert docx to pdf bytes =====
def _docx_to_pdf(file_bytes: bytes) -> bytes:
    """
    Try converting DOCX -> PDF for better layout fidelity with Layout Parser.
    Strategy:
      1) Aspose.Words (requires Aspose license)
      2) docx2pdf (Windows only - COM)
      3) Fallback: raise error (handled by caller)
    """
    # Option A: Aspose.Words (if available & licensed)
    try:
        import aspose.words as aw
        with tempfile.TemporaryDirectory() as tmpdir:
            in_path = os.path.join(tmpdir, "input.docx")
            out_path = os.path.join(tmpdir, "output.pdf")
            with open(in_path, "wb") as f_in:
                f_in.write(file_bytes)
            doc = aw.Document(in_path)
            doc.save(out_path)
            with open(out_path, "rb") as f_out:
                return f_out.read()
    except Exception:
        # swallow and try next option
        pass

    # Option B: docx2pdf (Windows only; may fail on Linux)
    try:
        from docx2pdf import convert  # type: ignore
        with tempfile.TemporaryDirectory() as tmpdir:
            in_path = os.path.join(tmpdir, "input.docx")
            out_path = os.path.join(tmpdir, "output.pdf")
            with open(in_path, "wb") as f_in:
                f_in.write(file_bytes)
            # convert may raise on non-Windows environments
            convert(in_path, out_path)
            with open(out_path, "rb") as f_out:
                return f_out.read()
    except Exception as e:
        raise RuntimeError(f"DOCX->PDF conversion failed (docx2pdf/Aspose not available or conversion error): {e}")

def _docx_text_fallback(file_bytes: bytes) -> str:
    # As a last resort, extract text without layout (for MVP continuity)
    try:
        import docx  # python-docx
        doc = docx.Document(BytesIO(file_bytes))
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        raise RuntimeError(f"DOCX text extraction failed: {e}")

def _text_from_layout(full_text: str, layout) -> str:
    """
    Extracts slices from full_text using the layout.text_anchor.text_segments.
    Returns empty string if no valid anchor present.
    """
    if not layout or not getattr(layout, "text_anchor", None) or not getattr(layout.text_anchor, "text_segments", None):
        return ""
    parts = []
    for seg in layout.text_anchor.text_segments:
        try:
            start = int(seg.start_index) if seg.start_index is not None else 0
            end = int(seg.end_index) if seg.end_index is not None else 0
        except Exception:
            # Defensive: if indices are not numeric, skip this segment
            continue
        parts.append(full_text[start:end])
    return "".join(parts)

def _cleanup_text(s: str) -> str:
    # de-hyphenate and normalize whitespace
    s = re.sub(r"(\w)-\n(\w)", r"\1\2", s)   # de-hyphenate at line breaks
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{2,}", "\n", s)
    return s.strip()

def _simple_paragraph_split(text: str):
    chunks = re.split(r"\n\s*\n|(?<=[.!?])\s+\n?", text)
    return [c.strip() for c in chunks if c.strip()]

def _simple_blocks(text: str):
    return [{"id": i, "text": t, "type": "paragraph", "page": 1} for i, t in enumerate(_simple_paragraph_split(text), 1)]

def _map_layout_to_blocks(doc: documentai.Document) -> Dict[str, Any]:
    """
    Prefer Layout Parser's chunked_document if present.
    Fallback to page paragraphs/blocks if needed.
    """
    full_text = doc.text or ""
    blocks = []
    id_counter = 1

    # 1) Use chunked_document chunks if available (ideal for rewriting)
    chunked = getattr(doc, "chunked_document", None)
    if chunked and getattr(chunked, "chunks", None):
        for ch in chunked.chunks:
            # Chunk objects typically expose `text` (not `layout`)
            ch_text = getattr(ch, "text", None) or getattr(ch, "content", None) or ""
            txt = _cleanup_text(ch_text)
            if txt:
                # try common page field names; fall back to 0
                page_val = getattr(ch, "page_ref", None) or getattr(ch, "page_number", None) or 0
                blocks.append({"id": id_counter, "text": txt, "type": "chunk", "page": page_val})
                id_counter += 1

    # 2) If no chunks, fall back to page paragraphs/blocks
    if not blocks and getattr(doc, "pages", None):
        for p_index, page in enumerate(doc.pages):
            para_list = getattr(page, "paragraphs", []) or getattr(page, "blocks", [])
            for para in para_list:
                txt = _text_from_layout(full_text, getattr(para, "layout", None))
                txt = _cleanup_text(txt)
                if txt:
                    blocks.append({"id": id_counter, "text": txt, "type": "paragraph", "page": p_index + 1})
                    id_counter += 1

    # 3) Last resort: split full text
    if not blocks and full_text:
        for t in _simple_paragraph_split(full_text):
            blocks.append({"id": id_counter, "text": t, "type": "paragraph", "page": 1})
            id_counter += 1

    return {"full_text": full_text, "blocks": blocks}

# --- Public entry point ---
def extract_text_and_blocks(file_bytes: bytes, filename: str, content_type: str | None) -> Dict[str, Any]:
    """
    Always uses the Layout Parser processor when sending documents to Document AI.
    - PDF/images: send directly.
    - DOCX: convert to PDF (preferred), else text fallback.
    - TXT: use as-is, no OCR.
    Returns dict with full_text and normalized blocks for the frontend.
    """
    ext = (os.path.splitext(filename)[1] or "").lower()
    mime = (content_type or "").lower()

    # Normalize MIME by extension when missing or generic
    if ext == ".pdf":
        mime = PDF_MIME
    elif ext == ".txt":
        mime = TXT_MIME
    elif ext == ".docx":
        mime = DOCX_MIME
    elif ext in [".jpg", ".jpeg"]:
        mime = "image/jpeg"
    elif ext == ".png":
        mime = "image/png"
    elif ext in [".tif", ".tiff"]:
        mime = "image/tiff"

    # TXT: bypass OCR
    if mime == TXT_MIME:
        text = file_bytes.decode("utf-8", errors="replace")
        return {"full_text": text, "blocks": _simple_blocks(text)}

    # DOCX: convert to PDF for best results with Layout Parser
    if mime == DOCX_MIME:
        try:
            pdf_bytes = _docx_to_pdf(file_bytes)
            doc = _process_with_layout(pdf_bytes, PDF_MIME)
            return _map_layout_to_blocks(doc)
        except Exception:
            # Fallback: naive DOCX text extraction
            text = _docx_text_fallback(file_bytes)
            return {"full_text": text, "blocks": _simple_blocks(text)}

    # PDFs and supported images: send directly
    if mime == PDF_MIME or mime in SUPPORTED_IMAGE_MIMES:
        doc = _process_with_layout(file_bytes, mime)
        return _map_layout_to_blocks(doc)

    # Unknown: try as PDF; if that fails, treat as text
    try:
        doc = _process_with_layout(file_bytes, PDF_MIME)
        return _map_layout_to_blocks(doc)
    except Exception:
        text = file_bytes.decode("utf-8", errors="replace")
        return {"full_text": text, "blocks": _simple_blocks(text)}

if __name__ == "__main__":
    print("Key path:", GOOGLE_APPLICATION_CREDENTIALS)
    print("Processor:", _processor_name())
    c = _client()
    # safer access to transport host; fallback if attribute names change
    host = getattr(getattr(c, "_transport", None), "_host", None) or getattr(getattr(c, "transport", None), "_host", None)
    print("Connected to:", host)
